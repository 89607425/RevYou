"""PRD parsing service: extract text from files, identify sections, recognize images."""
import re
import os
import uuid
from typing import Optional
from datetime import datetime, timezone

from app.core.config import settings


def parse_prd_structure(prd_text: str) -> dict:
    """Parse PRD text to extract section/chapter structure.

    Supports Chinese numbered chapters (一、/1./第一章 etc.) and Markdown headings.
    """
    sections = []
    lines = prd_text.split("\n")

    # Patterns for section headers
    patterns = [
        # Chinese chapter: 一、/ 第一章 / 第1章
        (r"^#*\s*(第[一二三四五六七八九十\d]+章)", 1),
        # Numbered section: 1.1 / 2.3.4
        (r"^#*\s*(\d+(?:\.\d+)*)\s", 1),
        # Chinese numbered: 一、/ 二、/ 三、
        (r"^#*\s*([一二三四五六七八九十]{1,3}、)", 1),
        # Markdown: ## / ### / ####
        (r"^(#{1,6})\s+([^#].+)$", 1),  # level = len(#)
    ]

    char_pos = 0
    for line_num, line in enumerate(lines):
        line_stripped = line.strip()
        char_pos += len(line) + 1  # +1 for newline

        for pattern, default_level in patterns:
            m = re.match(pattern, line_stripped)
            if m:
                if pattern.startswith(r"^(#{1,6})"):
                    level = len(m.group(1))
                    title = m.group(2).strip()
                else:
                    level = default_level
                    title = line_stripped

                section_id = f"S{len(sections) + 1}"
                sections.append({
                    "section_id": section_id,
                    "title": title[:100],
                    "level": level,
                    "char_range": [char_pos - len(line) - 1, char_pos],
                    "children": [],
                })
                break

    return {
        "sections": build_section_tree(sections),
        "total_sections": len(sections),
        "total_chars": len(prd_text),
    }


def build_section_tree(sections: list) -> list:
    """Build a hierarchical tree from flat section list."""
    root = []
    stack = []
    for s in sections:
        while stack and stack[-1]["level"] >= s["level"]:
            stack.pop()
        if stack:
            stack[-1]["children"].append(s)
        else:
            root.append(s)
        stack.append(s)
    return root


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file. Accepts file path or bytes."""
    try:
        import fitz
    except ImportError:
        return "[PDF解析需要安装PyMuPDF]"

    if isinstance(file_path, bytes):
        doc = fitz.open(stream=file_path, filetype="pdf")
    else:
        doc = fitz.open(file_path)
    text_content = ""
    for page_num, page in enumerate(doc):
        text_content += page.get_text()

    doc.close()
    return text_content


def parse_docx(file_path: str) -> str:
    """Extract text from a Word document. Accepts file path or bytes."""
    try:
        from docx import Document
        from io import BytesIO
    except ImportError:
        return "[Word解析需要安装python-docx]"

    if isinstance(file_path, bytes):
        doc = Document(BytesIO(file_path))
    else:
        doc = Document(file_path)
    text_content = ""
    for para in doc.paragraphs:
        text_content += para.text + "\n"
    return text_content


def enrich_prd_with_image_text(prd_text: str, prd_images: list) -> str:
    """Inject recognized image text into PRD content for LLM context."""
    completed_images = [img for img in prd_images if img.get("recognition_status") == "COMPLETED" and img.get("recognition_result")]
    if not completed_images:
        return prd_text

    enriched = prd_text + "\n\n---\n## 附件图片识别内容\n\n"
    for img in completed_images:
        enriched += f"### 图片: {img['filename']}\n{img['recognition_result']}\n\n"
    return enriched


def validate_file(filename: str, filesize: int) -> Optional[str]:
    """Validate uploaded file. Returns error message or None if valid."""
    allowed_exts = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp"}
    ext = os.path.splitext(filename)[1].lower()

    if ext not in allowed_exts:
        return f"不支持的文件格式: {ext}，支持: {', '.join(sorted(allowed_exts))}"

    is_image = ext in {".png", ".jpg", ".jpeg", ".webp"}
    max_size = 10 * 1024 * 1024 if is_image else settings.MAX_UPLOAD_SIZE

    if filesize > max_size:
        size_mb = max_size / 1024 / 1024
        return f"文件超过 {size_mb}MB 限制"

    return None
